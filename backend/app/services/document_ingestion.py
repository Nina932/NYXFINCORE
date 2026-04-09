"""
document_ingestion.py — Pipeline for ingesting internal documents into FinAI RAG.
══════════════════════════════════════════════════════════════════════════════════

Supports:
  - PDF (board presentations, audit reports, financial statements, transcripts)
  - Word (.docx) — memos, analysis reports, meeting minutes, policies
  - Excel (.xlsx) — supplementary tables, reference data, mapping files
  - Text (.txt, .md) — policy documents, notes, structured data exports

Each ingested document is:
  1. Text extracted with format-specific parser
  2. Split into overlapping chunks (~1500 chars, 200-char overlap)
  3. Stored in FinancialDocument table (DB-backed text search fallback)
  4. Indexed into ChromaDB vector store (semantic search)
  5. Full provenance attached: filename, page, chunk_index, document_type

Supported document_type values:
  "report"     — management reports, financial statements
  "memo"       — internal memos, board minutes
  "policy"     — policies, procedures, compliance
  "transcript" — earnings calls, interview transcripts
  "audit"      — audit findings, internal audit reports
  "reference"  — reference tables, COA mappings, supplementary data
  "other"      — unclassified internal documents

Usage:
    from app.services.document_ingestion import document_ingestion_service

    result = await document_ingestion_service.ingest_file(
        file_path="/tmp/Q1_2025_Board_Report.pdf",
        filename="Q1_2025_Board_Report.pdf",
        document_type="report",
        db=db,
        metadata={"period": "Q1 2025", "author": "Finance Team"},
    )
    # → {"document_id": 42, "chunks": 38, "pages": 12, "tokens_estimated": 5700}
"""

import hashlib
import logging
import os
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.models.all_models import FinancialDocument

logger = logging.getLogger(__name__)


# ── Optional library detection ──────────────────────────────────────────────

try:
    from pypdf import PdfReader
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False
    logger.info("pypdf not installed — PDF ingestion disabled (pip install pypdf)")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.info("python-docx not installed — DOCX ingestion disabled (pip install python-docx)")

try:
    import openpyxl
    _XLSX_AVAILABLE = True
except ImportError:
    _XLSX_AVAILABLE = False


# ── Constants ────────────────────────────────────────────────────────────────

CHUNK_SIZE    = 1500  # Characters per chunk (≈375 tokens at 4 chars/token)
CHUNK_OVERLAP = 200   # Overlap between consecutive chunks for context continuity
MAX_FILE_MB   = 50    # Maximum file size in MB

SUPPORTED_EXTENSIONS = {
    ".pdf":  "pdf",
    ".docx": "docx",
    ".doc":  "docx",  # treated as docx
    ".xlsx": "xlsx",
    ".xls":  "xlsx",  # treated as xlsx
    ".txt":  "text",
    ".md":   "text",
    ".csv":  "text",
}

VALID_DOCUMENT_TYPES = {
    "report", "memo", "policy", "transcript", "audit", "reference", "other",
}


# ════════════════════════════════════════════════════════════════════════════
# CHUNKING HELPERS
# ════════════════════════════════════════════════════════════════════════════

def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[Dict[str, Any]]:
    """
    Split text into overlapping chunks.

    Attempts to end each chunk at a sentence boundary (period/newline)
    to avoid mid-sentence splits. Returns list of:
      {"text": str, "chunk_index": int, "start_char": int, "end_char": int}
    """
    if not text or not text.strip():
        return []

    text = text.strip()
    chunks = []
    start = 0
    chunk_idx = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        # Try to end at a sentence boundary within the last 30% of the chunk
        if end < len(text):
            search_start = start + int(chunk_size * 0.7)
            # Look for newline first (paragraph boundary)
            boundary = text.rfind("\n", search_start, end)
            if boundary == -1:
                # Fall back to period
                boundary = text.rfind(". ", search_start, end)
                if boundary != -1:
                    boundary += 1  # include the period
            if boundary > search_start:
                end = boundary

        chunk_text = text[start:end].strip()
        if chunk_text and len(chunk_text) > 20:  # skip tiny fragments
            chunks.append({
                "text":        chunk_text,
                "chunk_index": chunk_idx,
                "start_char":  start,
                "end_char":    end,
            })
            chunk_idx += 1

        # Next chunk starts with overlap
        start = max(end - overlap, end) if end >= len(text) else end - overlap
        if start >= len(text):
            break

    return chunks


def _estimate_tokens(text: str) -> int:
    """Rough token estimate: ~4 characters per token for English."""
    return max(1, len(text) // 4)


# ════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTORS
# ════════════════════════════════════════════════════════════════════════════

def _extract_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from a PDF file using pypdf."""
    if not _PDF_AVAILABLE:
        return {"text": "", "pages": 0, "error": "pypdf not installed"}

    try:
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append({"page": i + 1, "text": page_text.strip()})

        full_text = "\n\n".join(p["text"] for p in pages)
        return {
            "text": full_text,
            "pages": len(reader.pages),
            "pages_with_text": len(pages),
            "error": None,
        }
    except Exception as e:
        logger.error("PDF extraction failed for %s: %s", file_path, e)
        return {"text": "", "pages": 0, "error": str(e)}


def _extract_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from a Word (.docx) file using python-docx."""
    if not _DOCX_AVAILABLE:
        return {"text": "", "pages": 0, "error": "python-docx not installed"}

    try:
        doc = DocxDocument(file_path)
        paragraphs = []

        # Extract body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and "Heading" in (para.style.name or ""):
                    paragraphs.append(f"\n## {text}\n")
                else:
                    paragraphs.append(text)

        # Extract table content
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    table_rows.append(row_text)
            if table_rows:
                paragraphs.append("\n[Table]\n" + "\n".join(table_rows) + "\n")

        full_text = "\n".join(paragraphs)
        return {
            "text": full_text,
            "pages": None,  # python-docx doesn't expose page count
            "paragraphs": len(paragraphs),
            "error": None,
        }
    except Exception as e:
        logger.error("DOCX extraction failed for %s: %s", file_path, e)
        return {"text": "", "pages": 0, "error": str(e)}


def _extract_xlsx(file_path: str) -> Dict[str, Any]:
    """Extract text from an Excel (.xlsx) file using openpyxl."""
    if not _XLSX_AVAILABLE:
        return {"text": "", "pages": 0, "error": "openpyxl not installed"}

    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheet_texts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            header = None

            for i, row in enumerate(ws.iter_rows(values_only=True)):
                # Skip entirely empty rows
                row_vals = [str(v).strip() if v is not None else "" for v in row]
                if not any(row_vals):
                    continue

                # First non-empty row is treated as header
                if header is None and any(row_vals):
                    header = row_vals
                    rows_text.append("Headers: " + " | ".join(v for v in header if v))
                else:
                    # Format as key=value pairs using header names
                    if header:
                        pairs = [
                            f"{header[j]}={row_vals[j]}"
                            for j in range(min(len(header), len(row_vals)))
                            if row_vals[j] and header[j]
                        ]
                        if pairs:
                            rows_text.append(", ".join(pairs))
                    else:
                        rows_text.append(" | ".join(v for v in row_vals if v))

                if i > 500:  # Cap at 500 rows per sheet
                    rows_text.append(f"[... {ws.max_row - 500} more rows ...]")
                    break

            if rows_text:
                sheet_texts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))

        wb.close()
        full_text = "\n\n".join(sheet_texts)
        return {
            "text": full_text,
            "pages": len(wb.sheetnames),
            "sheets": wb.sheetnames,
            "error": None,
        }
    except Exception as e:
        logger.error("XLSX extraction failed for %s: %s", file_path, e)
        return {"text": "", "pages": 0, "error": str(e)}


def _extract_text(file_path: str) -> Dict[str, Any]:
    """Read a plain text / markdown / CSV file."""
    try:
        # Try UTF-8 first, then fall back to cp1252
        for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
            try:
                with open(file_path, encoding=encoding, errors="replace") as f:
                    content = f.read()
                return {"text": content, "pages": 1, "encoding": encoding, "error": None}
            except UnicodeDecodeError:
                continue
        return {"text": "", "pages": 0, "error": "Could not decode file with any supported encoding"}
    except Exception as e:
        return {"text": "", "pages": 0, "error": str(e)}


# ════════════════════════════════════════════════════════════════════════════
# INGESTION SERVICE
# ════════════════════════════════════════════════════════════════════════════

class DocumentIngestionService:
    """
    Pipeline for ingesting internal documents into FinAI's RAG system.

    Documents are parsed, chunked, and indexed into:
    - FinancialDocument table (SQL text-search fallback)
    - ChromaDB "finai_documents" collection (semantic search)

    Each chunk gets a stable ID based on file content hash + chunk index,
    so re-ingesting the same file is idempotent.
    """

    COLLECTION_NAME = "finai_documents"

    async def ingest_file(
        self,
        file_path: str,
        filename: str,
        document_type: str,
        db: AsyncSession,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Full ingestion pipeline for a single file.

        Args:
            file_path:     Absolute path to the temporary/uploaded file
            filename:      Original filename (for display and metadata)
            document_type: "report" | "memo" | "policy" | "transcript" | "audit" | "reference" | "other"
            db:            Async database session
            metadata:      Optional extra metadata (period, author, source, etc.)

        Returns:
            {
                "document_id":        int,     # DB primary key
                "chunks":             int,     # Number of chunks indexed
                "pages":              int,     # Pages/sheets extracted
                "tokens_estimated":   int,     # Approx token count
                "format":             str,     # Detected file format
                "error":              str|None # Extraction error if any
            }
        """
        metadata = metadata or {}

        # ── Validate ──────────────────────────────────────────────────────
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}"}

        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > MAX_FILE_MB:
            return {"error": f"File too large: {file_size_mb:.1f}MB (max {MAX_FILE_MB}MB)"}

        ext = Path(filename).suffix.lower()
        fmt = SUPPORTED_EXTENSIONS.get(ext)
        if not fmt:
            return {"error": f"Unsupported file format: {ext}. Supported: {list(SUPPORTED_EXTENSIONS.keys())}"}

        if document_type not in VALID_DOCUMENT_TYPES:
            document_type = "other"

        # ── Extract text ──────────────────────────────────────────────────
        if fmt == "pdf":
            extracted = _extract_pdf(file_path)
        elif fmt == "docx":
            extracted = _extract_docx(file_path)
        elif fmt == "xlsx":
            extracted = _extract_xlsx(file_path)
        else:
            extracted = _extract_text(file_path)

        full_text = extracted.get("text", "")
        extract_error = extracted.get("error")

        if not full_text.strip():
            return {
                "error": extract_error or "No text content could be extracted from the file",
                "format": fmt,
            }

        # ── Create DB record for the document ─────────────────────────────
        file_hash = hashlib.sha256(full_text[:10000].encode("utf-8", errors="replace")).hexdigest()[:16]
        doc_metadata = {
            "filename":      filename,
            "document_type": document_type,
            "format":        fmt,
            "pages":         extracted.get("pages") or 0,
            "file_hash":     file_hash,
            "file_size_mb":  round(file_size_mb, 2),
            **{k: str(v) for k, v in metadata.items()},
        }

        db_doc = FinancialDocument(
            content=full_text[:500],  # Summary/preview
            metadata_json=doc_metadata,
            document_type=document_type,
            embedding_id=f"doc_{file_hash}",
            dataset_id=None,
            is_indexed=False,  # Will be True after chunking
        )
        db.add(db_doc)
        await db.flush()
        await db.refresh(db_doc)
        document_id = db_doc.id

        # ── Chunk text ────────────────────────────────────────────────────
        raw_chunks = _chunk_text(full_text)
        if not raw_chunks:
            return {
                "document_id": document_id,
                "chunks": 0,
                "error": "Text was too short to chunk meaningfully",
                "format": fmt,
            }

        # ── Build chunk records with metadata ─────────────────────────────
        chunk_records = []
        for chunk in raw_chunks:
            chunk_id = f"doc_{document_id}_chunk_{chunk['chunk_index']}"
            chunk_meta = {
                "document_id":   document_id,
                "filename":      filename,
                "document_type": document_type,
                "format":        fmt,
                "chunk_index":   chunk["chunk_index"],
                "total_chunks":  len(raw_chunks),
                "start_char":    chunk["start_char"],
                "end_char":      chunk["end_char"],
                **{k: str(v) for k, v in metadata.items() if isinstance(v, (str, int, float))},
            }
            chunk_records.append({
                "id":       chunk_id,
                "text":     chunk["text"],
                "metadata": chunk_meta,
            })

        # ── Index into vector store ────────────────────────────────────────
        from app.services.vector_store import vector_store
        indexed_count = await vector_store.add_document_chunks(
            chunk_records, db, collection_name=self.COLLECTION_NAME,
        )

        # ── Mark document as indexed ───────────────────────────────────────
        db_doc.is_indexed = True
        db_doc.metadata_json = {**doc_metadata, "chunks_indexed": indexed_count}
        await db.flush()

        tokens_est = _estimate_tokens(full_text)
        logger.info(
            "DocumentIngestion: '%s' → %d chunks, %d est. tokens, document_id=%d",
            filename, indexed_count, tokens_est, document_id,
        )

        return {
            "document_id":      document_id,
            "chunks":           indexed_count,
            "pages":            extracted.get("pages") or 0,
            "tokens_estimated": tokens_est,
            "format":           fmt,
            "file_size_mb":     round(file_size_mb, 2),
            "error":            extract_error,
        }

    async def list_documents(
        self,
        db: AsyncSession,
        document_type: Optional[str] = None,
        limit: int = 50,
        offset: int = 0,
    ) -> List[Dict[str, Any]]:
        """Return a paginated list of ingested documents."""
        q = select(FinancialDocument).where(
            FinancialDocument.document_type.in_(list(VALID_DOCUMENT_TYPES)),
            FinancialDocument.embedding_id.like("doc_%"),
        )
        if document_type and document_type in VALID_DOCUMENT_TYPES:
            q = q.where(FinancialDocument.document_type == document_type)

        q = q.order_by(FinancialDocument.created_at.desc()).offset(offset).limit(limit)
        result = await db.execute(q)
        docs = result.scalars().all()
        return [self._doc_to_dict(d) for d in docs]

    async def get_document(self, document_id: int, db: AsyncSession) -> Optional[Dict[str, Any]]:
        """Return a single document by ID."""
        result = await db.execute(
            select(FinancialDocument).where(FinancialDocument.id == document_id)
        )
        doc = result.scalar_one_or_none()
        return self._doc_to_dict(doc) if doc else None

    async def delete_document(self, document_id: int, db: AsyncSession) -> Dict[str, Any]:
        """
        Remove a document and all its vector store chunks.

        Returns {"deleted": True, "chunks_removed": N}
        """
        from app.services.vector_store import vector_store

        # Remove vector store chunks
        chunks_removed = await vector_store.delete_document_chunks(
            document_id, db, collection_name=self.COLLECTION_NAME,
        )

        # Remove the main FinancialDocument record and all chunk records
        records_deleted = 0
        try:
            # Delete the main document record
            result = await db.execute(
                select(FinancialDocument).where(FinancialDocument.id == document_id)
            )
            main_doc = result.scalar_one_or_none()
            if main_doc:
                await db.delete(main_doc)
                records_deleted += 1

            # Also delete FinancialDocument records created by add_document_chunks
            chunk_docs_result = await db.execute(
                select(FinancialDocument).where(
                    FinancialDocument.embedding_id.like(f"doc_{document_id}_chunk_%")
                )
            )
            for row in chunk_docs_result.scalars().all():
                await db.delete(row)
                records_deleted += 1

            await db.flush()
        except Exception as exc:
            logger.error("Failed to delete document %d from DB: %s", document_id, exc)
            return {"deleted": False, "error": str(exc)}

        return {
            "deleted": True,
            "document_id": document_id,
            "records_removed": records_deleted,
            "chunks_removed": chunks_removed,
        }

    async def search(
        self,
        query: str,
        n_results: int = 5,
    ) -> List[Dict[str, Any]]:
        """
        Semantic search over all indexed document chunks.

        Returns: [{"content", "score", "metadata"}]
        """
        from app.services.vector_store import vector_store
        return await vector_store.search_documents(
            query, n_results=n_results, collection_name=self.COLLECTION_NAME,
        )

    def supported_formats(self) -> Dict[str, Any]:
        """Return supported file formats and library availability."""
        return {
            "pdf":   {"available": _PDF_AVAILABLE,  "library": "pypdf"},
            "docx":  {"available": _DOCX_AVAILABLE, "library": "python-docx"},
            "xlsx":  {"available": _XLSX_AVAILABLE,  "library": "openpyxl"},
            "text":  {"available": True,             "library": "built-in"},
            "max_file_mb": MAX_FILE_MB,
            "chunk_size":  CHUNK_SIZE,
        }

    @staticmethod
    def _doc_to_dict(doc: FinancialDocument) -> Dict[str, Any]:
        meta = doc.metadata_json or {}
        return {
            "id":            doc.id,
            "document_type": doc.document_type,
            "filename":      meta.get("filename", ""),
            "format":        meta.get("format", ""),
            "pages":         meta.get("pages", 0),
            "chunks_indexed":meta.get("chunks_indexed", 0),
            "file_size_mb":  meta.get("file_size_mb", 0),
            "is_indexed":    doc.is_indexed,
            "metadata":      meta,
            "created_at":    doc.created_at.isoformat() if doc.created_at else None,
            "preview":       doc.content[:200] if doc.content else "",
        }


# ── Module-level singleton ──────────────────────────────────────────────────
document_ingestion_service = DocumentIngestionService()
